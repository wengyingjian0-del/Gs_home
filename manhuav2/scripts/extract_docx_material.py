from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document


def main() -> None:
    source = Path(sys.argv[1])
    target = Path(sys.argv[2])
    doc = Document(source)
    payload: dict[str, object] = {
        "source": str(source),
        "paragraphs": [
            {"style": p.style.name if p.style else "", "text": p.text}
            for p in doc.paragraphs
            if p.text.strip()
        ],
        "tables": [],
    }
    tables: list[object] = []
    for index, table in enumerate(doc.tables, start=1):
        tables.append({
            "index": index,
            "rows": [[cell.text for cell in row.cells] for row in table.rows],
        })
    payload["tables"] = tables
    target.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


if __name__ == "__main__":
    main()
