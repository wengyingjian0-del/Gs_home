from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parents[1] / "画芽_儿童漫画AI生图工具_PRD_V2.1.docx"

INK = "20304A"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "667085"
LIGHT = "F2F4F7"
PALE_BLUE = "EAF2F8"
PALE_GOLD = "FFF7E6"
PALE_RED = "FDECEC"
WHITE = "FFFFFF"
GRID = "D0D5DD"


def set_run_font(run, ascii_name="Calibri", east_asia="Microsoft YaHei", size=None, color=None, bold=None, italic=None):
    run.font.name = ascii_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), ascii_name)
    rfonts.set(qn("w:hAnsi"), ascii_name)
    rfonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_paragraph(p, fill):
    ppr = p._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)


def paragraph_border_bottom(p, color=BLUE, size=12, space=6):
    ppr = p._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)


def set_cell_shading(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tcpr = cell._tc.get_or_add_tcPr()
    mar = tcpr.find(qn("w:tcMar"))
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tcpr.append(mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblpr = table._tbl.tblPr
    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(total))
    tblw.set(qn("w:type"), "dxa")
    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:w"), "120")
    tblind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(width))
            tcw.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    set_run_font(normal.font, size=11) if False else None
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167
    code = styles.add_style("Code Block", 1)
    code.font.name = "Consolas"
    code.font.size = Pt(8.2)
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.18)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(8)
    code.paragraph_format.line_spacing = 1.0


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    run = paragraph.add_run(" 页")
    set_run_font(run, size=9, color=MUTED)


def setup_page(doc):
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
        hp = section.header.paragraphs[0]
        hp.text = "画芽 · 产品需求文档"
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_run_font(hp.runs[0], size=9, color=MUTED)
        add_page_field(section.footer.paragraphs[0])


def add_text(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True, color=INK)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r, color=INK)
    else:
        r = p.add_run(text)
        set_run_font(r, color=INK)
    return p


def add_bullet(doc, text, numbered=False):
    p = doc.add_paragraph(style="List Number" if numbered else "List Bullet")
    r = p.add_run(text)
    set_run_font(r, color=INK)
    return p


def add_callout(doc, label, text, fill=PALE_BLUE, color=DARK_BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(8)
    shade_paragraph(p, fill)
    r = p.add_run(f"{label}  ")
    set_run_font(r, bold=True, color=color)
    r = p.add_run(text)
    set_run_font(r, color=INK)
    return p


def add_table(doc, headers, rows, widths, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, size=font_size, bold=True, color=INK)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(str(value)) <= 10 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, color=INK)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_code(doc, code):
    p = doc.add_paragraph(style="Code Block")
    shade_paragraph(p, "F7F7F8")
    for idx, line in enumerate(code.splitlines()):
        if idx:
            p.add_run().add_break()
        r = p.add_run(line)
        set_run_font(r, ascii_name="Consolas", east_asia="Microsoft YaHei", size=8.2, color="303030")
    return p


def add_screenshot_note(doc, where, page, action, state, crop):
    add_callout(doc, "建议插图", f"插入位置：{where}。截图页面：{page}。操作：{action}。展示状态：{state}。建议裁剪：{crop}。", fill=PALE_GOLD, color="7A5A00")


def add_title_page(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run("PRODUCT REQUIREMENTS DOCUMENT")
    set_run_font(r, size=10, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("画芽 · 儿童漫画 AI 生图工具")
    set_run_font(r, size=25, bold=True, color=INK)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("当前实现版 PRD / V2.1")
    set_run_font(r, size=14, color=MUTED)
    for label, value in (
        ("文档日期", "2026 年 7 月 16 日"),
        ("产品形态", "手机端优先 H5 / PWA；含四格漫画工作室与候选管理页"),
        ("事实基线", "当前仓库代码、原始 V1.0 方案、12 条人工测试记录、V1→V2 修改"),
        ("验证状态", "构建通过；14/14 自动化检查通过；视觉与性能仍需人工回归"),
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"{label}：")
        set_run_font(r, bold=True, color=INK)
        r = p.add_run(value)
        set_run_font(r, color=INK)
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(10)
    rule.paragraph_format.space_after = Pt(14)
    paragraph_border_bottom(rule, BLUE, 14, 7)
    add_callout(doc, "阅读说明", "本文只把仓库中已经存在、能从代码或测试确认的能力写成“已实现”。只有设计稿、数据库结构或原始方案中出现但尚未接通的能力，均列为“未实现”或“待确认”。", PALE_BLUE)
    add_callout(doc, "密钥安全", "文档不记录任何真实 API Key。曾在对话或其他可见位置出现过的密钥应立即作废并轮换；生产密钥只能保存在服务端环境变量或密钥管理系统中。", PALE_RED, "9B1C1C")
    doc.add_page_break()


def build():
    doc = Document()
    configure_styles(doc)
    setup_page(doc)
    add_title_page(doc)

    doc.add_heading("1. 文档结论与产品定位", level=1)
    add_text(doc, "画芽是一款面向儿童及其家长的移动端 AI 漫画创作工具。儿童通过卡片选择、自由文字或浏览器语音输入，创建角色并生成单幅漫画场景，也可进入四格漫画工作室，让 AI 先规划四格脚本，再逐格生成画面。系统在服务端进行提示结构化、安全预检、两张初稿竞争、定向精修与按需补救，只向创作端返回最终审核通过的一张图片。")
    add_callout(doc, "当前真实状态", "项目已进入可运行的本地开发版本，不再是 V1.0 文档所述的“尚未开发”。但账号、短信验证、正式监护人认证、D1 数据持久化和生产部署资源仍未接通，因此当前不能视为可正式面向儿童上线的成品。")
    add_table(doc, ["维度", "当前结论"], [
        ["核心价值", "降低儿童表达门槛；复用角色；用对话持续改图；自动筛选较安全、较符合要求的结果"],
        ["主要使用者", "核心假设为中国大陆 8—12 岁儿童及家长；界面另提供 6—7 岁和 13—15 岁模式，但差异化规则尚未真正实现"],
        ["首发载体", "手机端优先 H5 / PWA；竖屏独立显示；本地开发地址为 localhost:3000"],
        ["当前成熟度", "功能原型 / 内部测试版"],
        ["模型", "万相 wan2.7-image 生图；qwen3-vl-plus 默认承担提示改写、候选视觉评审和四格策划，可由环境变量替换"],
    ], [1900, 7460])

    doc.add_heading("2. 产品背景与要解决的问题", level=1)
    add_text(doc, "儿童有丰富的故事想法，但传统绘画需要技巧，普通生图工具又要求用户理解提示词、模型参数和安全边界。对儿童而言，真正困难的不是“按下生成”，而是把零散想法补充完整、让同一主角在不同画面中仍可识别、理解为什么某些请求不能生成，以及在失败后知道下一步该做什么。")
    add_bullet(doc, "表达门槛：儿童往往只说“一个小女孩”，缺少地点、动作或可识别外观。")
    add_bullet(doc, "角色漂移：连续创作时，发型、服装、主色、背景实体可能随模型改变。")
    add_bullet(doc, "边界误判：普通描述可能因服务商内容检查产生误拦截，例如特定服饰或身份词。")
    add_bullet(doc, "安全与隐私：危险行为、真实人物模仿、联系方式和地址不能直接进入生图流程。")
    add_bullet(doc, "等待成本：多候选与逐格四格生成耗时明显，必须在质量和速度之间取舍。")
    add_bullet(doc, "管理边界：儿童需要简单操作，家长需要下载、删除和隐私控制，但当前尚无真实账号隔离。")

    doc.add_heading("3. 目标用户与使用场景", level=1)
    add_table(doc, ["用户", "核心诉求", "当前可完成", "当前限制"], [
        ["儿童创作者", "把故事快速变成漫画，不学习专业提示词", "创建角色、单图、纯场景、对话改图、四格漫画", "安全词库与模型仍可能误拦；视觉一致性需人工回归"],
        ["家长/监护人", "控制下载、查看或删除作品、保护隐私", "本机设置昵称/年龄模式、下载开关、清除本机数据", "无登录、无监护人认证、无多儿童账号隔离"],
        ["内部运营/测试", "查看每轮中间图、评分和模型评审理由", "使用服务端管理员口令进入 /admin，查看理解、初稿、精修、补救和最终结果", "必须接通 R2；当前为单口令访问，不是完整角色权限系统"],
    ], [1500, 2400, 2900, 2560], 8.7)
    add_text(doc, "核心场景包括：首次创建角色；复用已保存角色进入新场景；只生成无人物的环境；从历史作品继续对话修改；把较长故事灵感转成四格分镜；家长开启下载或清除本机数据；内部人员复核初稿、精修、补救及最终评审结果。")

    doc.add_heading("4. 产品目标与成功判断", level=1)
    add_table(doc, ["目标", "产品判断", "现状"], [
        ["低门槛", "预设不是必填；自由描述能被识别；缺项先提醒，用户可坚持生成", "已实现，自动化覆盖"],
        ["角色可复用", "可命名、保存、切换多个角色，并在新场景中使用", "已实现本机保存；视觉一致性待人工回归"],
        ["安全可解释", "危险、隐私、真人模仿和绕过规则的请求在生图前被阻止并给出可行动提示", "已实现基础规则；覆盖范围不等同专业审核服务"],
        ["多轮优化", "后台先生成 2 张初稿并择优，再生成 1 张定向精修；精修未胜出或评分不足时追加 1 张补救图", "已实现；阈值效果和质量提升幅度待离线标注集验证"],
        ["连续修改", "保留最近对话和所选历史图，只覆盖本轮冲突字段", "已实现结构与自动化检查；真实图像保持度待复测"],
        ["可上线", "有真实账号、监护人同意、数据隔离、管理端鉴权、生产存储与合规审查", "未达成"],
    ], [1450, 5000, 2910], 8.8)
    add_callout(doc, "待确认指标", "首次成功生成率、P50/P95 生成时长、候选评审与人工判断一致率、角色一致性合格线、误拦截率、四格完整成功率均尚无有效样本量，不能在本版承诺具体数值。", PALE_GOLD, "7A5A00")

    doc.add_heading("5. 功能范围与当前实现状态", level=1)
    add_table(doc, ["模块", "能力", "状态", "说明"], [
        ["角色", "五步角色向导、自由描述、命名、保存、切换、删除", "已实现", "保存于当前浏览器 localStorage；角色卡图片目前是公共展示素材，不是每个角色的专属生成参考图"],
        ["单图", "带主角场景、纯场景、快捷地点、语音/文字补充", "已实现", "语音依赖浏览器 SpeechRecognition；兼容范围待确认"],
        ["提示完整度", "从自由提示识别主角/发型/服装/地点/动作；缺项提醒；允许继续", "已实现", "继续生成只跳过完整度要求，不跳过安全规则"],
        ["生成", "结构化视觉描述、边界词中性化、2 张初稿、1 张精修、按需 1 张补救、逐轮视觉评审", "已实现", "没有模型配置时明确报错；多轮调用会增加等待时间和成本"],
        ["历史改图", "历史图引用、最近 12 条用户对话、修改意图推断", "已实现", "需要 R2 referenceKey 才能可靠保存与继续；分支树未可视化"],
        ["作品", "作品列表、打开、引用继续、当前作品下载", "部分实现", "收藏心形未持久化；家长作品页只管理当前作品，不是完整作品管理"],
        ["四格", "六类故事、2000 字灵感、AI 分镜、编辑台词/旁白、四格生成和 PNG 下载", "已实现", "逐格串行；必须有保存角色；稳定完成依赖 R2"],
        ["家长", "昵称、年龄模式、下载开关、隐私说明、清除本机数据", "部分实现", "无登录和真实家长验证；年龄模式尚未驱动不同安全规则"],
        ["管理", "按任务查看四阶段链路、中间图、最终图、评分、耗时和理由，并按状态筛选", "已实现", "仅 R2 环境有数据；需配置 ADMIN_ACCESS_TOKEN"],
        ["数据层", "D1 表结构、R2 绑定配置", "结构已建", "主用户流程未调用 D1；生产实例与迁移状态待确认"],
        ["分享/社区/支付", "公开分享、评论、私信、额度或支付", "未实现", "不在当前版本范围"],
    ], [1200, 3150, 1050, 3960], 8.1)

    doc.add_heading("6. 核心功能说明", level=1)
    doc.add_heading("6.1 角色创建与角色卡", level=2)
    add_text(doc, "用户可创建勇敢男孩、聪明女孩、机器伙伴或动物朋友，并依次选择发型、服装、性格和画风。每一步都可以跳过，也可在补充输入中用自己的话描述。完成后输入不超过 20 字的角色名，角色结构化属性保存在当前浏览器。")
    add_bullet(doc, "当前角色卡展示统一的角色构建插图，而不是该角色真实生成头像。")
    add_bullet(doc, "删除角色不会删除已生成作品；被删除角色与历史作品之间没有服务器侧完整性约束。")
    add_screenshot_note(doc, "角色创建功能说明之后", "角色页 / 角色创建向导", "先创建两个不同名字的角色，再返回角色列表", "同时显示两张角色卡、名称、属性摘要和“当前角色”标识", "保留手机壳内完整页面，裁去浏览器地址栏和桌面空白")

    doc.add_heading("6.2 单幅漫画与纯场景", level=2)
    add_text(doc, "带主角模式按地点、动作、情绪、额外元素、天气/时间五步采集信息；纯场景模式只保留地点、环境元素和天气/时间，后端提示明确要求画面不出现人物或拟人角色。首页支持校园、森林、海边、梦境快捷入口。")
    add_text(doc, "用户不必先点击所有预设。系统会检查自由描述是否已经包含角色、外貌/发型、服装、地点和动作；缺少信息时显示可补充项，同时提供“仍然按当前描述生成”。")
    add_screenshot_note(doc, "提示完整度说明之后", "场景向导最后一步", "只输入“一个小女孩”，点击生成", "红色提示框列出缺失项，并显示“仍然按当前描述生成”按钮", "裁到模式开关、自由输入框、提示框和继续按钮")

    doc.add_heading("6.3 安全预检与边界提示处理", level=2)
    add_text(doc, "服务端先对结构化字段和自由描述做确定性检查。当前覆盖危险/伤害、个人信息、真实人物模仿和绕过安全规则四类。命中后直接返回儿童友好提示，不调用生图。即使用户选择“仍然生成”，也只能跳过描述完整度检查，不能跳过安全检查。")
    add_text(doc, "对可能触发服务商误判、但本身可用于安全儿童插图的边界性表达，系统先改写为中性视觉描述。例如将特定身份、品牌或象征性服饰转换为原创卡通角色、普通颜色与形状；若万相仍返回内容检查错误，再用更严格的中性提示自动重试一次。")
    add_callout(doc, "能力边界", "当前规则是有限正则与模型评审的组合，不是完整的儿童内容安全产品。正式上线前仍需接入经评估的内容安全服务、监护人同意流程、审计和误拦截申诉机制。", PALE_RED, "9B1C1C")

    doc.add_heading("6.4 多轮优化生成与管理员复核", level=2)
    add_text(doc, "一次用户创作由 3 个固定阶段和 1 个按需阶段组成。第一阶段由视觉模型完成需求理解、安全判断和结构化 VisualSpec，不生成图片。第二阶段调用 wan2.7-image 生成 1536×2048 的两张初稿，视觉模型按意图、角色一致性、画面质量和儿童安全自动选出较好的一张。第三阶段以入选初稿为唯一参考，生成一张定向精修图，只改善评审指出的问题，并与初稿再次比较。")
    add_text(doc, "如果精修图没有胜过初稿，或入选结果的意图、画质、安全、角色分数仍低于补救门槛，系统进入第四阶段：基于当前最佳图生成一张最终温和补救图，再与当前最佳图比较。达到要求时提前结束，不强制生成补救图。儿童端始终只收到最终审核通过的一张图片。")
    add_text(doc, "基础入选门槛仍为：安全分不低于 80、意图分不低于 75，带主角时角色分不低于 70。触发补救使用更高的内部判断线：意图低于 82、画质低于 78、安全低于 85、角色低于 72，或精修没有被选中。管理员使用服务端 ADMIN_ACCESS_TOKEN 登录 /admin，可查看各轮中间图、四项评分、耗时、最终交付标记和评审理由。")
    add_text(doc, "在接通 R2 时，当前代码会保存本次任务的初稿、精修图及可能产生的补救图，并记录阶段关系。它们不会出现在儿童作品历史，但会进入管理员质量复核区。这与原始 V1“删除未入选候选”的要求不同，保留期限和自动删除策略仍待确认。")
    add_screenshot_note(doc, "多轮优化说明之后", "/admin 多轮优化管理台", "配置管理员口令并完成一次成功生图后登录", "展示需求理解、初稿竞争、定向精修、按需补救四阶段；标出最终交付图、每轮耗时、四项分数和评审备注", "裁取一个任务的完整优化链路，不显示管理员口令、地址栏、请求密钥或开发控制台")

    doc.add_heading("6.5 对话改图与历史版本", level=2)
    add_text(doc, "成功生成后进入对话改图。用户可说“换成雨天”“让她挥手”“开心一点”等，前端把最近最多 12 条用户指令连接为上下文，后端判断本轮修改的是背景、动作、表情或细节。引用历史图时，后端从 R2 读取该图和生成上下文，并要求未冲突字段继续保持。")
    add_bullet(doc, "每次只允许引用一张系统生成历史图。")
    add_bullet(doc, "参考图里的偶发背景角色、漂浮装饰和未明确要求的实体不应继承。")
    add_bullet(doc, "作品集以时间列表显示历史版本；没有可视化分支树，也没有真正的 parentVersion 关系写入 D1。")
    add_screenshot_note(doc, "多轮改图说明之后", "对话改图页", "基于历史图连续发送“换成雨天”与“让她挥手”", "对话中显示两轮要求、两张结果、当前参考版本和“从这个版本继续”", "保留完整对话线程，裁去底部系统任务栏")

    doc.add_heading("6.6 四格漫画工作室", level=2)
    add_text(doc, "用户选择一个已保存角色和六类故事之一，可输入最多 2000 字灵感。策划模型返回标题和四格脚本，每格包含场景、动作、情绪、背景、台词和旁白；用户可先改台词与旁白，再确认生成。四张成图按顺序生成，每一格都会执行两张初稿竞争、定向精修和按需补救，后一格引用前一格最终结果以保持主角连续。完成后可查看 2×2 或竖向版式，并下载整张 PNG。")
    add_callout(doc, "等待说明", "四格漫画需要连续完成 4 个单图任务，而每个单图任务现在包含两张初稿和至少一张精修，必要时还有一张补救图。因此整体耗时和模型成本会显著高于旧版，必须通过真实压测决定是否为四格单独采用更轻量的优化策略。", PALE_GOLD, "7A5A00")
    add_screenshot_note(doc, "四格功能说明之后", "/four-comic", "选角色与故事类型，输入一段灵感，先生成分镜，再完成四格图片", "上半部显示可编辑分镜，下半部显示完整 2×2 漫画册和下载按钮", "分别截分镜区和完整漫画册；不要把两块压缩成不可读的一张长图")

    doc.add_heading("6.7 家长设置与数据控制", level=2)
    add_text(doc, "家长空间当前提供儿童昵称、年龄模式、下载开关、当前作品查看/删除、隐私说明和本机数据清除。下载关闭时，儿童点击下载会收到“请家长开启”的提示。角色、昵称、设置、作品索引和对话记录主要存于 localStorage；图片和上下文在配置 R2 时由服务端保存。")
    add_callout(doc, "不是正式家长系统", "当前页面没有手机号登录、验证码、家长身份验证、儿童会话隔离或监护人同意记录。底部导航可直接进入设置，因此只能作为交互原型，不能按正式家长控制宣称。", PALE_RED, "9B1C1C")

    doc.add_heading("7. 用户流程", level=1)
    doc.add_heading("7.1 单图/纯场景主流程：文字节点与连接", level=2)
    add_text(doc, "流程从首页开始。用户选择已有角色、创建角色或进入纯场景；随后通过向导和自由输入形成描述。服务端先做安全检查，再判断描述是否完整：若缺项，返回提示，用户可以补充后再次提交，也可以明确继续；继续只绕过完整度检查。通过后，系统生成结构化视觉描述、两张初稿并择优，再产生一张定向精修图；只有精修未胜出或评分不足时才生成最后一张补救图。最终审核通过的一张进入对话改图和作品历史。")
    add_code(doc, """flowchart TD
    A[首页] --> B{创作方式}
    B -->|已有角色| C[选择角色]
    B -->|新角色| D[五步创建并命名保存]
    B -->|只画场景| E[纯场景三步向导]
    C --> F[场景向导与自由描述]
    D --> F
    E --> G[提交生成]
    F --> G
    G --> H{安全与隐私预检}
    H -->|阻止| I[解释原因与安全替代建议]
    H -->|通过| J{描述是否完整}
    J -->|缺项| K[提示缺失特点]
    K -->|补充| F
    K -->|仍然生成| L[使用中性默认设定]
    J -->|完整| M[结构化视觉描述]
    L --> M
    M --> N[生成2张初稿]
    N --> O{初稿视觉评审}
    O -->|无合格图| P[失败提示]
    O -->|选出1张| R[基于入选初稿生成1张定向精修]
    R --> S{精修是否胜出且达到补救门槛}
    S -->|是| Q[最终审核通过并交付1张]
    S -->|否| T[生成1张最终补救图]
    T --> U[与当前最佳图再次比较]
    U --> Q
    Q --> V[对话改图与作品历史]
    V -->|引用一张历史图| F""")
    add_text(doc, "英文生图提示词：")
    add_code(doc, """Create a clean, editable product workflow diagram for a Chinese mobile AI comic application named Huaya. Use a white background, navy text, blue process boxes, pale gold decision diamonds, and muted red blocked/error states. Show this exact left-to-right/top-to-bottom logic: Home -> choose existing character / create and name a new character / scene-only mode -> guided scene inputs and free text -> server safety and privacy pre-check -> if blocked, show child-friendly reason and safe alternatives; if allowed, check prompt completeness -> if incomplete, either add missing details or explicitly continue with neutral defaults -> convert to a structured visual specification -> generate exactly two initial Wan drafts -> evaluate and select the better safe draft -> generate exactly one targeted refinement using the selected draft as the only reference -> compare the refinement with the current best image -> if the refinement wins and meets the higher quality gate, deliver one final image; otherwise generate exactly one rescue image and compare again -> deliver only the final approved image -> conversational editing and history -> optionally reference one saved historical image and loop back. Use simple rounded rectangles, clear arrow labels, no decorative characters, no fictional screens, no extra steps, and reserve room for Chinese labels to be added later. 16:9 landscape, vector infographic, presentation-ready.""")

    doc.add_heading("7.2 四格漫画流程：文字节点与连接", level=2)
    add_text(doc, "四格流程先读取本机已保存角色。用户选择角色、故事类型和可选长文本灵感；策划模型生成四格脚本，用户可改台词和旁白。确认后从第 1 格到第 4 格顺序生成，每格内部走“需求理解、两张初稿、定向精修、按需补救”；第 2—4 格引用上一格最终结果。任一格失败即停止并保留当前页面错误信息；全部完成后才开放漫画册布局和整图下载。")
    add_code(doc, """flowchart TD
    A[进入四格漫画工作室] --> B{是否有已保存角色}
    B -->|否| C[返回角色页创建角色]
    B -->|是| D[选择角色、故事类型和灵感]
    D --> E[AI生成四格分镜]
    E --> F[用户编辑台词与旁白]
    F --> G[确认生成]
    G --> H[第1格：初稿竞争、精修、按需补救]
    H --> I[第2格：引用第1格最终图并多轮优化]
    I --> J[第3格：引用第2格最终图并多轮优化]
    J --> K[第4格：引用第3格最终图并多轮优化]
    H -->|失败| X[停止并显示具体错误]
    I -->|失败| X
    J -->|失败| X
    K -->|失败| X
    K -->|全部成功| L[选择2×2或竖向版式]
    L --> M[下载完整PNG]""")
    add_text(doc, "英文生图提示词：")
    add_code(doc, """Design a precise workflow infographic for the four-panel comic studio of a Chinese children's AI drawing app. White background, editorial product-document style, dark navy typography, coral accent for user actions, blue for AI actions, gold diamonds for decisions, red for errors. Exact sequence: open studio -> verify a locally saved character exists -> choose character, one of six story genres, and an optional story idea up to 2000 Chinese characters -> AI creates a four-panel script -> user edits dialogue and narration -> confirm -> for panel 1, generate two initial drafts, select one, create one targeted refinement, and create one rescue image only when the refinement does not meet the higher gate -> panel 2 references only panel 1's final image and repeats the same adaptive optimization -> repeat for panels 3 and 4 -> if any panel fails, stop and display a specific error -> when all four succeed, choose a 2-by-2 or vertical layout -> download one PNG. No parallel panel generation, no social sharing, no payment, no invented features. Clean vector diagram, 16:9, generous spacing, suitable for a PRD.""")

    doc.add_heading("8. 页面与交互说明", level=1)
    add_table(doc, ["页面/状态", "入口与主要操作", "关键反馈", "当前限制"], [
        ["首页", "角色入口、四格入口、快捷场景、只画场景、自由创作", "当前角色是否已准备", "产品名在 metadata 与页面文案不一致"],
        ["角色向导", "五步卡片；返回、跳过、自由输入、语音", "进度、选择态、语音失败提示", "预设跳过后仍需在结束时命名保存"],
        ["场景向导", "带主角/只画场景切换；选择、跳过、生成", "缺项提示、安全拦截、服务错误", "自由输入与预设的覆盖优先级仅由后端改写决定"],
        ["生成中", "等待", "说明 AI 正在理解、绘制两张初稿并定向优化", "无逐轮真实进度、取消、离页恢复"],
        ["对话改图", "输入新要求、使用建议、从历史版本继续", "当前参考版本、错误与继续按钮", "本地无 R2 时无法形成稳定 referenceKey"],
        ["结果页", "改表情/动作/背景/细节、收藏、下载、保存", "下载权限提示", "收藏不持久；“保存”与自动入历史的关系易混淆"],
        ["作品集", "打开历史图、引用继续", "最新/历史版本与时间", "没有收藏筛选、分支图和批量删除"],
        ["家长空间", "作品、身份、下载、隐私", "二次确认和本机清除", "无家长验证；作品管理不完整"],
        ["四格工作室", "选择角色/类型、策划、编辑、生成、布局、下载", "逐格状态和具体失败", "串行耗时长；无断点续传"],
        ["多轮优化管理", "登录后按任务查看理解、初稿、精修、补救、最终图、分数和耗时", "R2 未连接、口令未配置、旧版任务等状态", "当前只有单访问口令，无管理员角色分级和候选删除操作"],
    ], [1600, 3100, 2350, 2310], 8.0)
    add_screenshot_note(doc, "页面与交互章节开头", "首页", "准备一个已保存角色后回到首页", "显示角色已准备、四格入口、四个快捷场景和只画场景入口", "只保留手机端应用壳，宽度约 390 CSS 像素")
    add_screenshot_note(doc, "异常交互说明旁", "场景向导", "输入带有真实地址或危险动作的请求并提交", "分别截取隐私拒绝和危险内容拒绝状态，确保没有生成图", "裁取输入摘要、错误说明和返回/修改操作")
    add_screenshot_note(doc, "四格页面说明旁", "四格生成中", "确认脚本后开始生成", "显示“正在生成第 2 格，共 4 格”等真实阶段文案", "裁取分镜顶部、当前忙碌状态和前一格结果")

    doc.add_heading("9. 异常、边界与降级", level=1)
    add_table(doc, ["情况", "当前处理", "用户能否继续", "备注"], [
        ["描述不完整", "列出缺少的特点", "可补充或仍然生成", "只适用于无历史引用的新创作"],
        ["危险/伤害", "422 拒绝并给安全替代方向", "修改描述后重试", "有限规则集"],
        ["个人信息", "422 拒绝手机号、身份证、家庭/学校地址", "删除信息后重试", "真实姓名、普通学校名称等覆盖待加强"],
        ["真实人物模仿", "422 拒绝", "改成原创角色", "识别依赖有限关键词"],
        ["绕过安全规则", "422 拒绝", "直接描述安全情节", "“仍然生成”不能绕过"],
        ["模型未配置", "503 MODEL_NOT_CONFIGURED", "管理员配置新密钥", "不会调用外部接口或返回占位图"],
        ["万相内容检查", "将边界表达中性化后自动重试 1 次", "失败后调整词语", "返回请求 ID 便于排查"],
        ["两张初稿都不合格", "422 ALL_CANDIDATES_REJECTED，不进入精修", "重新调整描述", "不安全或明显跑题的初稿不会通过后续优化挽救"],
        ["R2 未连接", "本地返回供应商临时图 URL", "可查看，但无法稳定保存/引用", "URL 会过期"],
        ["下载未授权", "前端提示请家长开启", "家长开关后重试", "开关仅存于本机，可被技术用户修改"],
        ["四格某一格失败", "停止循环并显示第几格失败", "修正后重新生成整套", "无断点续传"],
        ["语音不可用", "提示直接打字", "是", "依赖浏览器实现"],
    ], [2100, 3750, 1450, 2060], 8.1)

    doc.add_heading("10. 测试结果与测试后的重要修改", level=1)
    add_text(doc, "原始人工测试集共 12 条：当时 7 条通过、5 条未通过，表面通过率约 58.3%。其中长文本用例只是“勉强通过”，若按稳健标准计入风险，则稳健通过为 5/12。当前已完成代码修复并新增 13 条自动化检查，构建和检查全部通过；两组结果的口径不同，不能合并为“当前 100% 通过”。")
    add_table(doc, ["#", "人工测试主题", "原结果", "V2 修改", "当前结论"], [
        ["1", "正常创建角色", "通过", "保留五步向导与保存", "基础流程通过"],
        ["2", "多个角色保存", "通过但信息弱", "增加角色命名、属性摘要、切换和删除", "功能已改；多人画面一致性待测"],
        ["3", "多轮上下文改图", "失败", "保存最近对话、推断修改字段、引用 R2 上下文", "自动化通过；视觉保持度待人工回归"],
        ["4", "基础场景/纯背景", "失败，误加角色", "新增纯场景模式；候选评审淘汰额外实体", "逻辑通过；真实图片待人工回归"],
        ["5", "引用历史图片", "通过", "强化单图 referenceKey 与历史入口", "依赖 R2 的环境回归"],
        ["6", "危险内容拦截", "失败", "生成前确定性拒绝，不能被继续按钮绕过", "自动化通过"],
        ["7", "隐私信息拦截", "通过", "保留并扩展错误码与儿童文案", "自动化通过"],
        ["8", "未授权下载", "通过", "保留家长下载开关和提示", "代码路径存在；无真实家长鉴权"],
        ["9", "约 2000 字长文本", "勉强通过、耗时长", "四格策划设 60 秒、1200 token；初稿 4→2，随后增加定向精修与按需补救", "性能仍待压测"],
        ["10", "信息过少", "失败，未澄清", "识别自由提示缺项，提醒后允许坚持生成", "自动化通过"],
        ["11", "矛盾要求，以最新为准", "通过", "明确最近对话优先，仅覆盖冲突字段", "逻辑已实现；视觉待回归"],
        ["12", "越狱/真实名人", "失败", "新增绕过规则与真人模仿前置拒绝", "自动化通过"],
    ], [420, 1800, 1500, 3400, 2240], 7.8)

    doc.add_heading("10.1 V1 → V2 修改总表", level=2)
    add_table(doc, ["改动领域", "V1/问题", "V2 当前实现", "验证"], [
        ["生成策略", "后台固定 4 张，等待较长", "改为 2 张初稿竞争 + 1 张定向精修 + 最多 1 张按需补救；管理员可见完整链路", "构建、类型与自动化"],
        ["幽灵/额外实体", "偶发多出白色漂浮角色", "提示只继承明确实体；评审对额外人物、动物、带脸装饰和幽灵状角色硬淘汰", "规则已实现，视觉待测"],
        ["纯场景", "无法生成无人物场景", "增加 scene-only 模式和三步流程", "自动化"],
        ["多轮上下文", "后续修改丢前文或多人物", "最近 12 条对话、结构化快照、修改意图推断", "自动化+待视觉"],
        ["角色管理", "多角色缺少名称和卡片信息", "可命名、切换、删除，显示属性摘要", "代码检查"],
        ["必填逻辑", "依赖预设选择", "从用户提示识别特点；缺项只提醒；允许继续", "自动化"],
        ["安全", "危险内容、越狱、真人模仿拦截不稳定", "服务端确定性前置拦截；继续按钮不可绕过", "自动化"],
        ["边界词", "安全描述也被万相 DataInspection 拦截", "结构化 VisualSpec、中性视觉词、检查失败后严格重试一次", "自动化模拟"],
        ["错误反馈", "只显示 400 或笼统失败", "区分内容检查、参数、密钥/权限并带请求 ID", "自动化"],
        ["长文本", "四格策划慢或不完整", "60 秒超时、1200 token、较低温度，灵感上限 2000 字", "功能实现，性能待测"],
        ["本地启动", "pnpm/Corepack 使用不稳定", "提供 start-dev.cmd、仓库内 pnpm.cmd 和构建脚本", "本机已运行"],
        ["运行时崩溃", "步骤索引导致读取 undefined.title", "按当前模式步骤长度夹紧索引", "自动化静态检查"],
    ], [1550, 2500, 3720, 1590], 7.8)

    doc.add_heading("11. 验收标准", level=1)
    add_text(doc, "以下标准按当前 V2 功能拆分。“通过”只代表可由当前代码或自动化确认；带“待人工”的项目必须在接通真实模型与 R2 的环境中验收。")
    add_table(doc, ["编号", "验收项", "通过条件", "状态"], [
        ["A01", "角色创建", "可跳过预设、命名并保存至少两个角色，切换后属性摘要正确", "可验收"],
        ["A02", "自由提示", "完整自由描述无需点击全部预设即可通过完整度检查", "自动化通过"],
        ["A03", "缺项提醒", "缺少关键特点时列出缺项；补充和“仍然生成”均可继续", "自动化通过"],
        ["A04", "安全不可绕过", "危险内容在 allowIncomplete=true 时仍被阻止", "自动化通过"],
        ["A05", "隐私/真人/越狱", "固定测试文本分别返回明确拒绝码与可行动提示", "自动化通过"],
        ["A06", "纯场景", "提示和最终图片均不出现未要求的人物、动物或拟人装饰", "逻辑通过/待人工"],
        ["A07", "多轮数量", "初稿请求 n=2；精修请求 n=1；只有精修未胜出或高门槛不足时补救请求 n=1", "代码通过"],
        ["A08", "逐轮择优", "基础入选满足安全≥80、意图≥75、带主角角色≥70；精修与补救均与当前最佳图重新比较", "代码通过/阈值待验证"],
        ["A09", "边界词", "示例服饰描述被转换为中性视觉字段；检查失败只自动重试一次", "自动化通过"],
        ["A10", "多轮改图", "最近指令优先；背景、动作、表情可被正确分类；未改字段尽量保持", "逻辑通过/待人工"],
        ["A11", "历史引用", "每次仅接受 generated/*.png 的一个 R2 历史引用", "代码通过/环境待验"],
        ["A12", "四格脚本", "输入≤2000 字后返回恰好四格，每格台词/旁白可编辑", "代码通过/模型待验"],
        ["A13", "四格成图", "4 格顺序完成，每格走初稿、精修和按需补救；下一格仅引用上一格最终图；可导出 2×2 和竖向 PNG", "待端到端"],
        ["A14", "下载权限", "关闭时阻止下载并提示家长开启；开启后触发 PNG 下载", "可验收"],
        ["A15", "密钥安全", "前端、仓库示例、文档和构建产物不含真实密钥", "自动化部分通过/需秘密扫描"],
        ["A16", "工程质量", "pnpm test 完成构建且 14/14 自动化检查通过", "已通过"],
    ], [700, 1850, 5150, 1660], 7.8)

    doc.add_heading("12. 技术与模型方案（通俗版）", level=1)
    add_text(doc, "前端是基于 React/Next 兼容技术栈的手机端网页，使用 Vite/Vinext 构建并适配 Cloudflare 运行环境。角色、儿童昵称、下载开关、作品索引和聊天消息目前主要保存在浏览器本机；图片生成与视觉评审由服务端接口完成，API Key 不下发到浏览器。")
    add_text(doc, "生成时，默认视觉模型 qwen3-vl-plus 先把用户的自然语言和历史要求整理为结构化 VisualSpec，包括主体、外观、服装、场景、动作、情绪、道具、光线、画风和构图。万相 wan2.7-image 按结构化描述生成两张初稿；视觉模型选出较好初稿后，万相以其为参考生成一张精修图。系统再次比较，必要时再产生一张补救图。四格工作室也用视觉模型生成四格脚本，并让每格调用同一套多轮优化。")
    add_table(doc, ["层", "当前实现", "边界"], [
        ["浏览器", "React 页面、localStorage、浏览器语音、Canvas 导出四格", "清缓存/换设备会丢本机数据；下载开关不是强安全边界"],
        ["服务端接口", "/api/generate、/api/four-comic/plan、/api/images/:id", "没有账号鉴权、限流、队列和任务恢复"],
        ["模型", "qwen3-vl-plus + wan2.7-image", "质量、时延、费用和区域权限受外部服务影响"],
        ["对象存储", "R2 保存初稿、精修、补救、最终图与上下文", "生产桶与生命周期策略待确认；当前会保留管理员可见的过程图"],
        ["数据库", "D1 schema 已定义", "主流程未调用；不能据此宣称账号和作品已持久化"],
    ], [1500, 4000, 3860], 8.4)

    doc.add_heading("13. 已知问题", level=1)
    add_table(doc, ["优先级", "问题", "影响", "建议"], [
        ["P0", "曾暴露的 API Key", "可能被滥用并产生费用", "立即作废并轮换；执行仓库和构建产物秘密扫描"],
        ["P0", "家长鉴权缺失；管理端仅单口令", "儿童可进入家长设置；管理员缺少账号、角色、审计和口令轮换界面", "上线前完成真实家长登录，并把管理端升级为账号、角色权限和审计系统"],
        ["P0", "D1 未接入主流程", "账号、作品和权限不能跨设备可靠保存", "实现服务端身份与数据读写，再做迁移和隔离测试"],
        ["P0", "多轮过程图全量保留", "初稿、精修和补救会增加隐私暴露面与存储成本，并与原 V1 删除未入选图目标冲突", "确定合规保留期；管理员复核窗口结束后自动删除非最终图"],
        ["P1", "角色卡没有专属参考图", "所谓角色一致性主要依赖文字字段，能力弱于原方案", "保存正式角色参考图并在新场景中引用"],
        ["P1", "本地无 R2 时 referenceKey 为空", "作品历史和四格连续生成不稳定", "开发环境也绑定可持久 R2 或提供受控本地对象存储"],
        ["P1", "收藏状态未持久化", "心形仅当前页面有效", "把 favorite 纳入作品对象并保存"],
        ["P1", "家长作品管理只看当前作品", "无法完成“全部作品管理”文案承诺", "改为完整列表、单项/批量删除和结果反馈"],
        ["P1", "生成中无离页恢复", "长任务中离开即失去可见状态", "引入任务 ID、队列、轮询/恢复和幂等"],
        ["P1", "四格无断点续传", "第 4 格失败可能需要重做全部", "保存逐格进度并支持从失败格重试"],
        ["P1", "产品命名不一致", "manifest 为“画芽”，metadata 为“漫小星”", "统一产品名、标题、应用名与文案"],
        ["P2", "输入长度口径不一致", "后端总长度 2500，但提示要求 2000；聊天仅 240", "在 PRD 和界面统一各入口上限"],
        ["P2", "年龄模式只改显示值", "用户以为安全规则已差异化", "实现真实规则或明确标注仅为资料设置"],
        ["P2", "未做正式无障碍与真机测试", "儿童、弱网和不同浏览器体验未知", "按目标机型执行 WCAG/触控/语音/下载测试"],
    ], [700, 2800, 2860, 3000], 7.8)

    doc.add_heading("14. 后续规划", level=1)
    add_table(doc, ["阶段", "目标", "主要内容", "进入条件"], [
        ["P0 上线门槛", "从原型变为受控测试版", "密钥轮换；家长/管理端鉴权；D1 接入；R2 生命周期；监护人同意；限流和审计", "安全、隐私和数据评审通过"],
        ["P1 质量闭环", "证明生成质量", "复跑 12 条人工集；扩展边界词、纯场景、额外实体、多轮一致性标注集；冻结评审阈值", "有足够样本和人工一致性结论"],
        ["P1 任务可靠性", "减少长等待损失", "异步任务、离页恢复、四格断点续传、失败格重试、进度与耗时统计", "真实模型压测完成"],
        ["P1 作品体系", "完成版本和家长管理", "真实人物参考图、父子版本、收藏、完整作品列表、删除级联与恢复策略", "D1/R2 关系稳定"],
        ["P2 体验优化", "提升儿童易用性", "统一品牌、真机语音、弱网、PWA 安装、无障碍、年龄文案", "可用性测试完成"],
        ["暂不规划", "控制范围", "社区、公开分享、评论私信、支付、视频和真人照片上传", "新的产品决策前不进入开发"],
    ], [1500, 2200, 4200, 1460], 8.2)

    doc.add_heading("15. 截图与配图执行清单", level=1)
    add_text(doc, "本 PRD 正文可以独立阅读，不依赖图片。若用于评审汇报，建议按下表补充真实产品截图；截图前必须隐藏地址栏、请求 ID、控制台、API Key、账号和任何真实个人信息。")
    add_table(doc, ["编号", "插入章节", "页面与操作", "要展示的状态", "裁剪建议"], [
        ["P01", "5/8", "首页，已有角色", "核心入口和信息架构", "390px 手机壳完整首屏"],
        ["P02", "6.1", "角色列表，建立两个名称不同角色", "多角色、摘要、当前角色", "页面标题到两个角色卡"],
        ["P03", "6.2", "场景向导，描述不完整", "缺项提醒与仍然生成", "输入区和提示框"],
        ["P04", "6.3/9", "提交危险或隐私内容", "明确阻止且无生成图", "错误卡和返回操作"],
        ["P05", "6.4", "/admin 登录后打开一次新任务", "四阶段时间线、两张初稿、精修图、可选补救图、最终交付标记、评分与耗时", "单个任务的完整优化链路"],
        ["P06", "6.5", "历史图连续改背景和动作", "多轮对话与当前参考版本", "对话线程主体"],
        ["P07", "6.6", "四格脚本生成后", "可编辑四格分镜", "分镜标题和四个脚本卡"],
        ["P08", "6.6", "四格全部完成", "2×2 漫画册与下载", "完整漫画册，不含桌面空白"],
        ["P09", "8", "四格生成中", "第 N 格/共 4 格状态", "状态文案与已完成格"],
        ["P10", "6.7", "家长设置", "下载开关、身份、隐私入口", "手机壳内完整设置页"],
    ], [600, 1100, 3000, 2500, 2160], 8.0)
    add_callout(doc, "其他配图", "除上述两张流程图外，不建议额外使用 AI 生成产品图。PRD 的事实性更适合用真实截图；装饰性插画会增加篇幅，却不能证明功能已经实现。", PALE_BLUE)

    doc.add_heading("16. 待确认事项", level=1)
    add_bullet(doc, "目标年龄是否正式锁定 8—12 岁；当前页面同时出现 6—7、8—12、13—15 三档。")
    add_bullet(doc, "正式产品名称采用“画芽”还是“漫小星”；当前代码存在不一致。")
    add_bullet(doc, "生产环境是否已经创建并绑定真实 D1/R2；仓库只有绑定名称和占位配置。")
    add_bullet(doc, "未入选候选图是否允许保留、保留多久、谁可查看，以及合规依据。")
    add_bullet(doc, "视觉评审模型、四项阈值和边界中性化规则是否通过足量人工标注验证。")
    add_bullet(doc, "单图和四格的目标耗时、失败率、日限额和成本预算。")
    add_bullet(doc, "家长登录、短信服务、监护人同意、数据删除时限和中国大陆儿童个人信息专项合规方案。")
    add_bullet(doc, "角色参考图的生成与存储方案；当前角色卡只有结构化文字和公共素材。")
    add_bullet(doc, "PWA 首批支持的手机浏览器、语音输入、下载和后台恢复兼容清单。")

    doc.add_heading("附录 A：材料与事实来源", level=1)
    add_bullet(doc, "原始方案：儿童漫画AI生图工具_MVP产品方案与PRD_v1.0.docx（2026-07-13，方案基线）。")
    add_bullet(doc, "当前实现：app、lib、db、tests、README、manifest、Cloudflare 绑定配置。")
    add_bullet(doc, "人工测试记录：用户提供的 12 条测试结果截图（正例、负例、边界与对抗）。")
    add_bullet(doc, "V1→V2 修改：本轮仓库实际代码变更和界面反馈。")
    add_bullet(doc, "验证：2026-07-16 执行 pnpm test，完成构建并通过 14/14 自动化检查；TypeScript 与本次相关 ESLint 检查通过。")
    add_callout(doc, "版本声明", "本文是当前真实实现的产品 PRD，不是上线合规证明，也不替代真实模型回归、儿童用户研究、安全评估或法律审查。", PALE_GOLD, "7A5A00")

    doc.core_properties.title = "画芽 · 儿童漫画 AI 生图工具 PRD V2.1"
    doc.core_properties.subject = "基于当前真实实现、测试结果与修改记录的产品需求文档"
    doc.core_properties.author = "产品团队"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
